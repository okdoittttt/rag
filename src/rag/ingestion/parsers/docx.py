"""Word(.docx) 파서

python-docx를 사용하여 .docx 파일에서 문단과 표 텍스트를 추출합니다.
"""

from __future__ import annotations

from pathlib import Path

from rag.ingestion.parsers.base import DocumentParser
from rag.logger import get_logger


logger = get_logger(__name__)


class DocxParser(DocumentParser):
    """Word(.docx) 파일 파서"""

    extensions = [".docx"]

    def parse(self, path: Path) -> str:
        """.docx 파일에서 텍스트를 추출한다.

        문단 텍스트를 순서대로 모으고, 표(table)는 각 행의 셀을 ``" | "``로
        연결하여 본문에 포함한다. ``python-docx`` 는 호출 시점에 지연 import 하여
        선택적 의존성으로 다룬다.

        Args:
            path: 파싱할 .docx 파일 경로.

        Returns:
            추출된 텍스트. 추출 실패 시 빈 문자열.

        Raises:
            ValueError: ``python-docx`` 가 설치되어 있지 않을 때 발생.
        """
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            logger.warning("docx_dependency_missing", error=str(e))
            raise ValueError(
                "python-docx가 설치되어 있지 않습니다. 'uv sync'로 의존성을 설치하세요."
            ) from e

        try:
            document = DocxDocument(str(path))
            parts: list[str] = []

            for para in document.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    line = " | ".join(c for c in cells if c)
                    if line:
                        parts.append(line)

            return "\n\n".join(parts)
        except Exception as e:
            logger.warning("docx_parse_failed", path=str(path), error=str(e))
            return ""
